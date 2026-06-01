from __future__ import annotations

import html
import re
import textwrap
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


EXPORT_CONTENT_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "odt": "application/vnd.oasis.opendocument.text",
    "archimate": "application/xml",
}

EXPORT_EXTENSIONS = {
    "pdf": "pdf",
    "docx": "docx",
    "odt": "odt",
    "archimate": "archimate.xml",
}

_FONT_NAME = "DejaVuSans"
_FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/local/share/fonts/DejaVuSans.ttf",
]

_ARCHIMATE_TYPE_MAP = {
    "business_actor": "BusinessActor",
    "business_role": "BusinessRole",
    "business_process": "BusinessProcess",
    "business_function": "BusinessFunction",
    "business_service": "BusinessService",
    "business_object": "BusinessObject",
    "data_object": "DataObject",
    "application_component": "ApplicationComponent",
    "application_service": "ApplicationService",
    "application_interface": "ApplicationInterface",
    "application_function": "ApplicationFunction",
    "node": "Node",
    "device": "Device",
    "system_software": "SystemSoftware",
    "technology_service": "TechnologyService",
    "technology_interface": "TechnologyInterface",
    "artifact": "Artifact",
}


def export_solution_document(solution: dict[str, Any], export_format: str) -> bytes:
    document = _solution_document(solution)
    return _export_document(document, export_format)


def export_protocol_document(protocol: dict[str, Any], export_format: str) -> bytes:
    document = _protocol_document(protocol)
    return _export_document(document, export_format)


def export_solution_archimate(solution: dict[str, Any]) -> bytes:
    model = solution.get("architecture_model") or {}
    entities = list(model.get("entities") or [])
    relations = list(model.get("relations") or [])

    ET.register_namespace("", "http://www.opengroup.org/xsd/archimate/3.0/")
    ET.register_namespace("xsi", "http://www.w3.org/2001/XMLSchema-instance")
    root = ET.Element(
        "model",
        {
            "xmlns": "http://www.opengroup.org/xsd/archimate/3.0/",
            "xmlns:xsi": "http://www.w3.org/2001/XMLSchema-instance",
            "identifier": _xml_id("model", solution.get("solution_version_id")),
        },
    )
    ET.SubElement(root, "name").text = str(solution.get("solution_title") or "Architecture model")

    elements_node = ET.SubElement(root, "elements")
    known_entity_ids: set[str] = set()
    for index, entity in enumerate(entities, start=1):
        entity_id = str(entity.get("entity_id") or entity.get("component_name") or f"entity-{index}")
        known_entity_ids.add(entity_id)
        element = ET.SubElement(
            elements_node,
            "element",
            {
                "identifier": _xml_id("element", entity_id),
                "{http://www.w3.org/2001/XMLSchema-instance}type": _archimate_type(entity),
            },
        )
        ET.SubElement(element, "name").text = str(entity.get("name") or entity.get("component_name") or entity_id)
        documentation = _clean_text(
            entity.get("description")
            or entity.get("role_description")
            or entity.get("technology_stack")
            or ""
        )
        if documentation:
            ET.SubElement(element, "documentation").text = documentation

    relationships_node = ET.SubElement(root, "relationships")
    for index, relation in enumerate(relations, start=1):
        source_id = str(relation.get("source_entity_id") or relation.get("source") or "")
        target_id = str(relation.get("target_entity_id") or relation.get("target") or "")
        if source_id not in known_entity_ids or target_id not in known_entity_ids:
            continue
        relationship = ET.SubElement(
            relationships_node,
            "relationship",
            {
                "identifier": _xml_id("relationship", relation.get("relation_id") or index),
                "{http://www.w3.org/2001/XMLSchema-instance}type": "AssociationRelationship",
                "source": _xml_id("element", source_id),
                "target": _xml_id("element", target_id),
            },
        )
        relation_name = str(relation.get("relation_type") or relation.get("description") or "").strip()
        if relation_name:
            ET.SubElement(relationship, "name").text = relation_name

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def export_filename(base_name: str, export_format: str) -> str:
    extension = EXPORT_EXTENSIONS.get(export_format, export_format)
    return f"{_slugify(base_name)}.{extension}"


def _export_document(document: list[tuple[str, str]], export_format: str) -> bytes:
    if export_format == "pdf":
        return _document_to_pdf(document)
    if export_format == "docx":
        return _document_to_docx(document)
    if export_format == "odt":
        return _document_to_odt(document)
    raise ValueError(f"Unsupported export format: {export_format}")


def _solution_document(solution: dict[str, Any]) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = [
        ("h1", str(solution.get("solution_title") or "Архитектурное решение")),
        ("p", str(solution.get("executive_summary") or "")),
    ]
    for section in solution.get("sections") or []:
        rows.append(("h2", str(section.get("title") or section.get("section_code") or "Раздел")))
        rows.append(("p", str(section.get("body_markdown") or "")))
    components = solution.get("components") or []
    if components:
        rows.append(("h2", "Компоненты"))
        for component in components:
            rows.append(("p", _join_fields(component, ["component_name", "role_description", "technology_stack"])))
    integrations = solution.get("integrations") or []
    if integrations:
        rows.append(("h2", "Интеграции"))
        for integration in integrations:
            rows.append(("p", _join_fields(integration, ["source_component", "target_component", "integration_type", "description"])))
    risks = solution.get("risks") or []
    if risks:
        rows.append(("h2", "Риски"))
        for risk in risks:
            rows.append(("p", _join_fields(risk, ["title", "severity", "description", "mitigation"])))
    return rows


def _protocol_document(protocol: dict[str, Any]) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = [
        ("h1", f"Протокол проверки {protocol.get('protocol_id') or ''}".strip()),
        ("p", f"Итоговый статус: {protocol.get('summary_status') or '—'}"),
        ("p", str(protocol.get("summary_text") or "")),
    ]
    basis_documents = protocol.get("basis_documents") or []
    if basis_documents:
        rows.append(("h2", "Документы-основания"))
        for item in basis_documents:
            rows.append(("p", _join_fields(item, ["title", "role_code", "version_ref"])))
    findings = protocol.get("findings") or []
    if findings:
        rows.append(("h2", "Проверки и замечания"))
        for finding in findings:
            rows.append(
                (
                    "p",
                    _join_fields(
                        finding,
                        ["rule_id", "rule_name", "status", "severity", "finding_text", "evidence", "related_section_ref"],
                    ),
                )
            )
    return rows


def _document_to_docx(rows: list[tuple[str, str]]) -> bytes:
    document = Document()
    for kind, text in rows:
        clean = _clean_text(text)
        if not clean:
            continue
        if kind == "h1":
            document.add_heading(clean, level=1)
        elif kind == "h2":
            document.add_heading(clean, level=2)
        else:
            for paragraph in clean.splitlines() or [clean]:
                document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _document_to_odt(rows: list[tuple[str, str]]) -> bytes:
    body_parts = []
    for kind, text in rows:
        clean = _clean_text(text)
        if not clean:
            continue
        tag = "text:h" if kind in {"h1", "h2"} else "text:p"
        attrs = ' text:outline-level="1"' if kind == "h1" else ' text:outline-level="2"' if kind == "h2" else ""
        for paragraph in clean.splitlines() or [clean]:
            body_parts.append(f"<{tag}{attrs}>{html.escape(paragraph)}</{tag}>")
    content_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" office:version="1.2">'
        "<office:body><office:text>"
        + "".join(body_parts)
        + "</office:text></office:body></office:document-content>"
    )
    manifest_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" '
        'manifest:version="1.2">'
        '<manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/>'
        '<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>'
        "</manifest:manifest>"
    )
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("mimetype", "application/vnd.oasis.opendocument.text", compress_type=zipfile.ZIP_STORED)
        archive.writestr("content.xml", content_xml)
        archive.writestr("META-INF/manifest.xml", manifest_xml)
    return buffer.getvalue()


def _document_to_pdf(rows: list[tuple[str, str]]) -> bytes:
    _register_pdf_font()
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 42
    y = height - margin

    def draw_line(text: str, *, font_size: int = 10, leading: int = 14) -> None:
        nonlocal y
        if y < margin + leading:
            pdf.showPage()
            pdf.setFont(_FONT_NAME, font_size)
            y = height - margin
        pdf.setFont(_FONT_NAME, font_size)
        pdf.drawString(margin, y, text)
        y -= leading

    for kind, text in rows:
        clean = _clean_text(text)
        if not clean:
            continue
        font_size = 16 if kind == "h1" else 13 if kind == "h2" else 10
        leading = 22 if kind == "h1" else 18 if kind == "h2" else 14
        for raw_line in clean.splitlines() or [clean]:
            wrapped = textwrap.wrap(raw_line, width=92) or [""]
            for line in wrapped:
                draw_line(line, font_size=font_size, leading=leading)
            if kind == "p":
                y -= 2
        if kind in {"h1", "h2"}:
            y -= 4
    pdf.save()
    return buffer.getvalue()


def _register_pdf_font() -> None:
    if _FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return
    for candidate in _FONT_CANDIDATES:
        if Path(candidate).exists():
            pdfmetrics.registerFont(TTFont(_FONT_NAME, candidate))
            return
    raise RuntimeError("DejaVuSans.ttf is required for PDF export")


def _archimate_type(entity: dict[str, Any]) -> str:
    code = str(entity.get("archimate_element_code") or "").strip().lower()
    if code in _ARCHIMATE_TYPE_MAP:
        return _ARCHIMATE_TYPE_MAP[code]
    layer = str(entity.get("archimate_layer") or entity.get("boundary_type") or "").lower()
    if "business" in layer:
        return "BusinessObject"
    if "data" in layer:
        return "DataObject"
    if "technology" in layer:
        return "Node"
    return "ApplicationComponent"


def _join_fields(value: Any, keys: list[str]) -> str:
    if not isinstance(value, dict):
        return _clean_text(value)
    parts = [str(value.get(key)).strip() for key in keys if value.get(key) not in (None, "")]
    return " — ".join(parts)


def _clean_text(value: Any) -> str:
    return re.sub(r"\n{3,}", "\n\n", str(value or "").replace("\r\n", "\n")).strip()


def _slugify(value: str) -> str:
    slug = re.sub(r"[^0-9A-Za-zА-Яа-я._-]+", "_", value.strip())
    return slug.strip("._-")[:120] or "artifact"


def _xml_id(prefix: str, value: Any) -> str:
    raw = re.sub(r"[^0-9A-Za-z_.-]+", "-", str(value or "").strip())
    return f"{prefix}-{raw.strip('-') or 'item'}"
