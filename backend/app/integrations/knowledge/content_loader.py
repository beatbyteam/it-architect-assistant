from __future__ import annotations

import hashlib
import io
import json
import re
from contextlib import suppress
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.document import Document as DocxNativeDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from pypdf import PdfReader

from app.integrations.knowledge.extraction_markers import (
    DOMAIN_SIGNAL_MARKERS,
    LEGAL_DISCLAIMER_MARKERS,
)
from app.integrations.knowledge.local_paths import is_local_path_reference, resolve_local_path
from app.integrations.knowledge.source_security import (
    assert_safe_remote_url,
    enforce_document_size_limit,
)

try:
    import trafilatura  # type: ignore[import-untyped]
except Exception:  # pragma: no cover
    trafilatura = None


class ContentLoadError(RuntimeError):
    pass


HEADING_RE = re.compile(r"^(#{1,6}\s+.+|\d+(?:\.\d+)*\s+.+)$")


@dataclass(slots=True)
class StructuredSection:
    heading: str | None
    content: str
    source_location: str | None = None
    level: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class NormalizedDocument:
    text: str
    content_format: str
    parser_name: str
    sections: list[StructuredSection] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


SUPPORTED_DOCUMENT_SUFFIXES = {
    ".pdf",
    ".docx",
    ".html",
    ".htm",
    ".md",
    ".markdown",
    ".txt",
    ".text",
    ".json",
    ".xlsx",
}

MEDIA_TYPE_SUFFIXES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/xhtml+xml": ".html",
    "text/html": ".html",
    "text/markdown": ".md",
    "text/x-markdown": ".md",
    "text/plain": ".txt",
    "application/json": ".json",
    "text/json": ".json",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
}


def fetch_uri(
    uri: str, timeout_sec: float = 30.0, max_size_bytes: int | None = None
) -> tuple[bytes, str, str | None]:
    try:
        if _is_local_path(uri):
            path = resolve_local_path(uri)
            if max_size_bytes is not None:
                enforce_document_size_limit(path.stat().st_size, max_size_bytes=max_size_bytes)
            data = path.read_bytes()
            return data, str(path), _guess_media_type(path.suffix.lower())

        parsed = urlparse(uri)
        if parsed.scheme == "file":
            path = resolve_local_path(uri)
            if max_size_bytes is not None:
                enforce_document_size_limit(path.stat().st_size, max_size_bytes=max_size_bytes)
            data = path.read_bytes()
            return data, str(path), _guess_media_type(path.suffix.lower())

        assert_safe_remote_url(uri)
        with (
            httpx.Client(timeout=timeout_sec, follow_redirects=True) as client,
            client.stream("GET", uri) as response,
        ):
            response.raise_for_status()
            assert_safe_remote_url(str(response.url))
            media_type = response.headers.get("content-type")
            chunks: list[bytes] = []
            total_bytes = 0
            for part in response.iter_bytes():
                if not part:
                    continue
                total_bytes += len(part)
                if max_size_bytes is not None:
                    enforce_document_size_limit(total_bytes, max_size_bytes=max_size_bytes)
                chunks.append(part)
            return b"".join(chunks), str(response.url), media_type
    except Exception as exc:  # pragma: no cover
        raise ContentLoadError(str(exc)) from exc


def normalize_document(uri: str, data: bytes, *, media_type: str | None = None) -> str:
    return normalize_document_payload(uri, data, media_type=media_type).text


def normalize_document_payload(
    uri: str, data: bytes, *, media_type: str | None = None
) -> NormalizedDocument:
    suffix = _detect_document_suffix(uri, media_type=media_type)
    if suffix == ".pdf":
        return _extract_pdf_payload(uri, data)
    if suffix == ".docx":
        return _extract_docx_payload(data)
    if suffix in {".html", ".htm"}:
        return _extract_html_payload(data)
    if suffix in {".md", ".markdown"}:
        return _extract_markdown_payload(data)
    if suffix in {".txt", ".text", ".json"}:
        return _extract_text_payload(data, suffix=suffix)
    if suffix == ".xlsx":
        return _extract_xlsx_payload(data)
    return _extract_fallback_payload(uri, data, suffix=suffix, media_type=media_type)


def checksum_sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def compute_embedding_key(
    content: str,
    *,
    profile_code: str | None = None,
    model_id: str | None = None,
    dimensions: int | None = None,
    task_mode: str = "document",
    template_version: str = "v1",
    normalization_version: str = "v1",
) -> str:
    payload = "\n".join(
        [
            f"profile={profile_code or 'default'}",
            f"model={model_id or 'default'}",
            f"dimensions={dimensions or 0}",
            f"task_mode={task_mode}",
            f"template_version={template_version}",
            f"normalization_version={normalization_version}",
            content,
        ]
    )
    return f"sha256:{hashlib.sha256(payload.encode('utf-8')).hexdigest()}"


def _is_local_path(uri: str) -> bool:
    return is_local_path_reference(uri)


def _detect_document_suffix(uri: str, *, media_type: str | None = None) -> str:
    suffix = _detect_suffix(uri)
    if suffix in SUPPORTED_DOCUMENT_SUFFIXES:
        return suffix
    return _detect_suffix_from_media_type(media_type)


def _detect_suffix_from_media_type(media_type: str | None) -> str:
    normalized = (media_type or "").split(";", 1)[0].strip().lower()
    return MEDIA_TYPE_SUFFIXES.get(normalized, "")


def _detect_suffix(uri: str) -> str:
    if _is_local_path(uri):
        try:
            return resolve_local_path(uri).suffix.lower()
        except Exception:
            return Path(uri).suffix.lower()
    parsed = urlparse(uri)
    return Path(parsed.path or uri).suffix.lower()


def _extract_fallback_payload(
    uri: str,
    data: bytes,
    *,
    suffix: str,
    media_type: str | None = None,
) -> NormalizedDocument:
    decoded = _decode_best_effort_text(data)
    parser_name = "text-fallback"
    content_format = "fallback_text"
    if decoded:
        text = decoded
    else:
        parser_name = "binary-fallback"
        content_format = "fallback_binary"
        filename = Path(urlparse(uri).path or uri).name or "document"
        suffix_label = suffix or "unknown"
        media_type_label = _normalize_media_type(media_type) or "unknown"
        text = (
            f"Binary document fallback\n\n"
            f"File: {filename}\n"
            f"Extension: {suffix_label}\n"
            f"Media type: {media_type_label}\n"
            f"Size bytes: {len(data)}\n"
            "The original format is not natively supported, so structured text extraction "
            "was not available."
        )
    return NormalizedDocument(
        text=text,
        content_format=content_format,
        parser_name=parser_name,
        sections=[StructuredSection(heading=None, content=text, source_location="fallback:1")],
        metadata={
            "fallback_used": True,
            "detected_suffix": suffix or None,
            "detected_media_type": _normalize_media_type(media_type),
            "canonical_text_stats": _text_stats(text),
        },
    )


def _decode_best_effort_text(data: bytes) -> str:
    candidates: list[tuple[float, str]] = []
    for encoding in ("utf-8", "utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "cp1251"):
        with suppress(UnicodeDecodeError, LookupError):
            decoded = data.decode(encoding)
            if decoded and decoded.count("\x00") / len(decoded) > 0.05:
                continue
            normalized = decoded.replace("\x00", "").strip()
            if not normalized:
                continue
            total = max(len(normalized), 1)
            printable = sum(1 for ch in normalized if ch.isprintable() or ch in "\r\n\t")
            alpha_num = sum(1 for ch in normalized if ch.isalnum())
            score = (printable / total) + min(alpha_num / total, 0.4)
            if printable / total < 0.85:
                continue
            candidates.append((score, normalized))
    if not candidates:
        return ""
    return max(candidates, key=lambda item: item[0])[1]


def _normalize_media_type(media_type: str | None) -> str | None:
    normalized = (media_type or "").split(";", 1)[0].strip().lower()
    return normalized or None


def _extract_pdf_payload(uri: str, data: bytes) -> NormalizedDocument:
    try:
        reader = PdfReader(io.BytesIO(data))
        sections: list[StructuredSection] = []
        page_texts: list[str] = []
        skipped_pages: list[dict[str, Any]] = []
        seen_page_signatures: set[str] = set()
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if not page_text:
                skipped_pages.append({"page_number": index, "reason": "empty_page"})
                continue
            skip_reason = _classify_skippable_pdf_page(
                page_text, seen_signatures=seen_page_signatures
            )
            if skip_reason:
                skipped_pages.append({"page_number": index, "reason": skip_reason})
                continue
            page_texts.append(page_text)
            sections.append(
                StructuredSection(
                    heading=f"Page {index}",
                    content=page_text,
                    source_location=f"page:{index}",
                    level=1,
                    metadata={"page_number": index},
                )
            )
        text = "\n\n".join(
            _render_section(section) for section in sections if section.content.strip()
        )
        if not text.strip():
            fallback = _extract_fallback_payload(uri, data, suffix=".pdf", media_type="application/pdf")
            fallback.metadata = {
                **fallback.metadata,
                "pdf_fallback_reason": "no_extractable_text",
                "skipped_pages": skipped_pages,
            }
            return fallback
        page_map = _build_pdf_page_map([section for section in sections if section.content.strip()])
        return NormalizedDocument(
            text=text,
            content_format="pdf",
            parser_name="pypdf",
            sections=sections,
            metadata={
                "page_count": len(reader.pages),
                "section_count": len(sections),
                "skipped_page_count": len(skipped_pages),
                "skipped_pages": skipped_pages,
                "page_map": page_map,
                "canonical_text_stats": _text_stats(text),
                "parser_warnings": [],
            },
        )
    except Exception as exc:
        fallback = _extract_fallback_payload(uri, data, suffix=".pdf", media_type="application/pdf")
        fallback.metadata = {
            **fallback.metadata,
            "pdf_fallback_reason": "parse_failed",
            "pdf_parse_error": str(exc),
        }
        return fallback


def _extract_docx_payload(data: bytes) -> NormalizedDocument:
    try:
        document = DocxDocument(io.BytesIO(data))
        sections: list[StructuredSection] = []
        current_heading: str | None = None
        current_heading_level: int | None = None
        buffer: list[str] = []
        section_index = 0
        table_index = 0
        block_start = 1
        block_end = 0
        current_block_index = 0

        def flush() -> None:
            nonlocal \
                section_index, \
                buffer, \
                current_heading, \
                current_heading_level, \
                block_start, \
                block_end
            content = "\n".join(line for line in buffer if line.strip()).strip()
            if not content:
                buffer = []
                return
            section_index += 1
            start_index = block_start
            end_index = max(block_end, start_index)
            heading_slug = _slugify_source_fragment(current_heading) if current_heading else None
            location = f"docx:block:{start_index}-{end_index}"
            if heading_slug:
                location = f"{location};heading:{heading_slug}"
            sections.append(
                StructuredSection(
                    heading=current_heading,
                    content=content,
                    source_location=location,
                    level=current_heading_level,
                    metadata={"heading": current_heading, "block_range": [start_index, end_index]},
                )
            )
            buffer = []
            block_start = current_block_index + 1

        for block in _iter_docx_blocks(document):
            current_block_index += 1
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not text:
                    continue
                style_name = getattr(getattr(block, "style", None), "name", "") or ""
                if style_name.lower().startswith("heading"):
                    flush()
                    current_heading = text
                    current_heading_level = _parse_docx_heading_level(style_name)
                    block_start = current_block_index
                    block_end = current_block_index
                else:
                    buffer.append(text)
                    block_end = current_block_index
                continue
            if isinstance(block, Table):
                table_index += 1
                rows: list[str] = []
                for row in block.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        rows.append(" | ".join(values))
                if rows:
                    buffer.append(f"Table {table_index}\n" + "\n".join(rows))
                    block_end = current_block_index
        flush()
        if not sections:
            flat = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            sections = [
                StructuredSection(
                    heading=None, content="\n\n".join(flat), source_location="section:1"
                )
            ]
        canonical_text = "\n\n".join(
            _render_section(section) for section in sections if section.content.strip()
        )
        return NormalizedDocument(
            text=canonical_text,
            content_format="docx",
            parser_name="python-docx",
            sections=sections,
            metadata={
                "section_count": len(sections),
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "section_map": _build_section_map(sections),
                "canonical_text_stats": _text_stats(canonical_text),
                "parser_warnings": [],
            },
        )
    except Exception as exc:  # pragma: no cover
        raise ContentLoadError(f"Failed to parse DOCX: {exc}") from exc


def _iter_docx_blocks(document: DocxNativeDocument):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _extract_xlsx_payload(data: bytes) -> NormalizedDocument:
    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheet_count = len(workbook.sheetnames)
        skipped_sheets: list[dict[str, Any]] = []
        sections: list[StructuredSection] = []
        parser_warnings: list[str] = []
        total_non_empty_rows = 0
        total_non_empty_cells = 0

        for sheet_index, worksheet in enumerate(workbook.worksheets, start=1):
            if _should_skip_xlsx_sheet(worksheet.title, workbook.sheetnames):
                skipped_sheets.append(
                    {
                        "sheet_name": worksheet.title,
                        "sheet_index": sheet_index,
                        "reason": "duplicate_or_low_value_sheet",
                    }
                )
                continue
            rows: list[str] = []
            first_row: int | None = None
            last_row: int | None = None
            non_empty_cells = 0
            max_column = 0

            for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                values = [_normalize_spreadsheet_cell(value) for value in row]
                while values and not values[-1]:
                    values.pop()
                if not any(values):
                    continue
                if first_row is None:
                    first_row = row_index
                last_row = row_index
                max_column = max(max_column, len(values))
                non_empty_cells += sum(1 for value in values if value)
                rows.append(" | ".join(_escape_table_cell(value) for value in values))

            if not rows:
                continue

            total_non_empty_rows += len(rows)
            total_non_empty_cells += non_empty_cells
            row_range = [first_row or 1, last_row or first_row or 1]
            heading = f"Sheet: {worksheet.title}"
            sections.append(
                StructuredSection(
                    heading=heading,
                    content="\n".join(rows),
                    source_location=(
                        f"xlsx:sheet:{sheet_index};rows:{row_range[0]}-{row_range[1]};"
                        f"sheet:{_slugify_source_fragment(worksheet.title)}"
                    ),
                    level=1,
                    metadata={
                        "sheet_name": worksheet.title,
                        "sheet_index": sheet_index,
                        "row_range": row_range,
                        "non_empty_row_count": len(rows),
                        "non_empty_cell_count": non_empty_cells,
                        "max_column": max_column,
                    },
                )
            )

        workbook.close()

        if not sections:
            parser_warnings.append("workbook_has_no_extractable_cells")

        canonical_text = "\n\n".join(
            _render_section(section) for section in sections if section.content.strip()
        )
        return NormalizedDocument(
            text=canonical_text,
            content_format="xlsx",
            parser_name="openpyxl",
            sections=sections,
            metadata={
                "sheet_count": sheet_count,
                "section_count": len(sections),
                "non_empty_row_count": total_non_empty_rows,
                "non_empty_cell_count": total_non_empty_cells,
                "skipped_sheets": skipped_sheets,
                "section_map": _build_section_map(sections),
                "canonical_text_stats": _text_stats(canonical_text),
                "parser_warnings": parser_warnings,
            },
        )
    except Exception as exc:  # pragma: no cover
        raise ContentLoadError(f"Failed to parse XLSX: {exc}") from exc


def _should_skip_xlsx_sheet(sheet_name: str, all_sheet_names: list[str]) -> bool:
    normalized = re.sub(r"\s+", " ", sheet_name.strip().casefold())
    if normalized in {"document title", "notice", "administrative appendix"}:
        return True
    non_summary_sheets = [
        name
        for name in all_sheet_names
        if re.sub(r"\s+", " ", name.strip().casefold())
        not in {"document title", "notice", "administrative appendix", "all domains"}
    ]
    return normalized == "all domains" and len(non_summary_sheets) >= 3


def _normalize_spreadsheet_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return re.sub(r"\s+", " ", value).strip()
    return str(value).strip()


def _escape_table_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\r", " ").replace("\n", " ").strip()


def _extract_html_payload(data: bytes) -> NormalizedDocument:
    html_text = data.decode("utf-8", errors="ignore")
    extracted = None
    if trafilatura is not None:
        extracted = trafilatura.extract(
            html_text, include_comments=False, include_tables=True, include_links=True
        )
    soup = BeautifulSoup(html_text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    sections: list[StructuredSection] = []
    current_heading: str | None = None
    current_level: int | None = None
    buffer: list[str] = []
    section_index = 0
    block_start = 1
    block_end = 0
    block_index = 0

    def flush() -> None:
        nonlocal section_index, buffer, current_heading, current_level, block_start, block_end
        content = "\n".join(line for line in buffer if line.strip()).strip()
        if not content:
            buffer = []
            return
        section_index += 1
        start_index = block_start
        end_index = max(block_end, start_index)
        heading_slug = _slugify_source_fragment(current_heading) if current_heading else None
        location = f"html:block:{start_index}-{end_index}"
        if heading_slug:
            location = f"{location};heading:{heading_slug}"
        sections.append(
            StructuredSection(
                heading=current_heading,
                content=content,
                source_location=location,
                level=current_level,
                metadata={"heading": current_heading, "block_range": [start_index, end_index]},
            )
        )
        buffer = []
        block_start = block_index + 1

    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text_value = tag.get_text(" ", strip=True)
        if not text_value:
            continue
        block_index += 1
        if tag.name.startswith("h"):
            flush()
            current_heading = text_value
            current_level = int(tag.name[1:]) if tag.name[1:].isdigit() else 1
            block_start = block_index
            block_end = block_index
        else:
            buffer.append(text_value)
            block_end = block_index
    flush()

    text = extracted or "\n\n".join(
        _render_section(section) for section in sections if section.content.strip()
    )
    if not sections and text.strip():
        sections = [
            StructuredSection(
                heading=None,
                content=text.strip(),
                source_location="html:block:1-1",
                metadata={"block_range": [1, 1]},
            )
        ]
    canonical_text = text.strip()
    return NormalizedDocument(
        text=canonical_text,
        content_format="html",
        parser_name="trafilatura" if extracted else "beautifulsoup4",
        sections=sections,
        metadata={
            "section_count": len(sections),
            "title": soup.title.get_text(strip=True) if soup.title else None,
            "section_map": _build_section_map(sections),
            "canonical_text_stats": _text_stats(canonical_text),
            "parser_warnings": [],
        },
    )


def _extract_markdown_payload(data: bytes) -> NormalizedDocument:
    text = data.decode("utf-8", errors="ignore")
    sections = _split_markdown_sections(text)
    canonical_text = "\n\n".join(
        _render_section(section) for section in sections if section.content.strip()
    )
    return NormalizedDocument(
        text=canonical_text,
        content_format="markdown",
        parser_name="markdown-structured",
        sections=sections,
        metadata={
            "section_count": len(sections),
            "section_map": _build_section_map(sections),
            "canonical_text_stats": _text_stats(canonical_text),
            "parser_warnings": [],
        },
    )


def _extract_text_payload(data: bytes, *, suffix: str) -> NormalizedDocument:
    text = data.decode("utf-8", errors="ignore")
    if suffix == ".json":
        with suppress(json.JSONDecodeError):
            text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
    sections = _split_markdown_sections(text)
    parser = "json-normalized" if suffix == ".json" else "plain-text"
    content_format = suffix.lstrip(".") or "txt"
    canonical_text = "\n\n".join(
        _render_section(section) for section in sections if section.content.strip()
    )
    return NormalizedDocument(
        text=canonical_text,
        content_format=content_format,
        parser_name=parser,
        sections=sections,
        metadata={
            "section_count": len(sections),
            "section_map": _build_section_map(sections),
            "canonical_text_stats": _text_stats(canonical_text),
            "parser_warnings": [],
        },
    )


def _split_markdown_sections(text: str) -> list[StructuredSection]:
    lines = text.splitlines()
    sections: list[StructuredSection] = []
    heading: str | None = None
    level: int | None = None
    buffer: list[str] = []
    index = 0
    start_line = 1
    end_line = 0
    current_line_no = 0

    def flush() -> None:
        nonlocal index, buffer, heading, level, start_line, end_line
        content = "\n".join(buffer).strip()
        if not content and heading is None:
            buffer = []
            return
        index += 1
        first_line = start_line
        last_line = max(end_line, first_line)
        heading_slug = _slugify_source_fragment(heading) if heading else None
        location = f"lines:{first_line}-{last_line}"
        if heading_slug:
            location = f"{location};heading:{heading_slug}"
        sections.append(
            StructuredSection(
                heading=heading,
                content=content or (heading or ""),
                source_location=location,
                level=level,
                metadata={
                    "heading": heading,
                    "level": level,
                    "line_range": [first_line, last_line],
                },
            )
        )
        buffer = []
        start_line = current_line_no + 1

    for raw_line in lines:
        current_line_no += 1
        line = raw_line.rstrip()
        if HEADING_RE.match(line.strip()):
            flush()
            stripped = line.strip()
            start_line = current_line_no
            end_line = current_line_no
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                heading = stripped.lstrip("#").strip()
            else:
                level = 1
                heading = stripped
        else:
            buffer.append(line)
            end_line = current_line_no
    flush()
    if not sections and text.strip():
        sections = [
            StructuredSection(
                heading=None,
                content=text.strip(),
                source_location=f"lines:1-{max(1, len(lines))}",
                metadata={"line_range": [1, max(1, len(lines))]},
            )
        ]
    return sections


def _classify_skippable_pdf_page(
    text: str, *, seen_signatures: set[str]
) -> str | None:
    compact = re.sub(r"\s+", " ", text or "").strip()
    if not compact:
        return "empty_page"
    tokens = re.findall(r"\w+", compact, flags=re.UNICODE)
    if len(tokens) <= 4 and not any(any(ch.isalpha() for ch in token) for token in tokens):
        return "low_information_page"
    signature = _canonical_page_signature(compact)
    if signature in seen_signatures:
        return "duplicate_page"
    seen_signatures.add(signature)
    if _looks_like_table_of_contents_page(text):
        return "table_of_contents"
    if _looks_like_legal_disclaimer(compact):
        return "legal_disclaimer"
    return None


def _canonical_page_signature(text: str) -> str:
    normalized = re.sub(r"\W+", " ", text.lower(), flags=re.UNICODE)
    normalized = re.sub(r"\b\d{1,4}\b", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return hashlib.sha1(normalized[:4000].encode("utf-8")).hexdigest()


def _looks_like_legal_disclaimer(text: str) -> bool:
    lowered = text.lower()
    legal_hits = sum(1 for marker in LEGAL_DISCLAIMER_MARKERS if marker in lowered)
    token_count = len(re.findall(r"\w+", lowered, flags=re.UNICODE))
    evaluation_footer = all(
        marker in lowered
        for marker in ("all rights reserved", "evaluation copy", "not for redistribution")
    )
    if evaluation_footer and token_count <= 90:
        return True
    if legal_hits < 2 and not any(
        marker in lowered for marker in {"legal disclaimer", "отказ от ответственности"}
    ):
        return False
    domain_hits = sum(1 for marker in DOMAIN_SIGNAL_MARKERS if marker in lowered)
    legal_density = legal_hits / max(token_count, 1)
    return domain_hits == 0 or legal_hits >= 5 or legal_density >= 0.08


def _looks_like_table_of_contents_page(text: str) -> bool:
    lowered = text.lower()
    has_toc_title = "table of contents" in lowered or "содержание" in lowered
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) < 6:
        return False
    toc_like_lines = 0
    for line in lines:
        normalized = re.sub(r"\s+", " ", line)
        if re.search(r"\.{3,}", normalized):
            toc_like_lines += 1
        elif re.search(r"\s(?:[ivxlcdm]+|\d{1,4})$", normalized, flags=re.IGNORECASE):
            toc_like_lines += 1
    toc_ratio = toc_like_lines / max(len(lines), 1)
    if has_toc_title:
        return toc_like_lines >= 5 or toc_ratio >= 0.45
    return toc_like_lines >= 8 and toc_ratio >= 0.55


def _build_pdf_page_map(sections: list[StructuredSection]) -> list[dict[str, Any]]:
    page_map: list[dict[str, Any]] = []
    offset = 0
    for index, section in enumerate(sections, start=1):
        page_text = _render_section(section)
        start_offset = offset
        end_offset = start_offset + len(page_text)
        page_number = int((section.metadata or {}).get("page_number") or index)
        page_map.append(
            {
                "page_number": page_number,
                "source_location": section.source_location or f"page:{page_number}",
                "char_range": [start_offset, end_offset],
                "approx_tokens": max(0, len(page_text.split())),
            }
        )
        offset = end_offset + 2
    return page_map


def _build_section_map(sections: list[StructuredSection]) -> list[dict[str, Any]]:
    mapping: list[dict[str, Any]] = []
    offset = 0
    for index, section in enumerate(sections, start=1):
        rendered = _render_section(section)
        start_offset = offset
        end_offset = start_offset + len(rendered)
        item = {
            "index": index,
            "heading": section.heading,
            "source_location": section.source_location,
            "level": section.level,
            "char_range": [start_offset, end_offset],
            "metadata": dict(section.metadata or {}),
        }
        mapping.append(item)
        offset = end_offset + 2
    return mapping


def _text_stats(text: str) -> dict[str, int]:
    stripped = text or ""
    return {
        "char_count": len(stripped),
        "line_count": len(stripped.splitlines()),
        "word_count": len(stripped.split()),
    }


def _parse_docx_heading_level(style_name: str) -> int | None:
    match = re.search(r"(\d+)$", style_name.strip())
    if not match:
        return 1 if style_name.strip() else None
    try:
        return int(match.group(1))
    except ValueError:
        return 1


def _slugify_source_fragment(value: str | None) -> str:
    text = re.sub(r"[^a-zA-Z0-9а-яА-ЯёЁ]+", "-", (value or "").strip().lower()).strip("-")
    return text[:80] if text else "section"


def _render_section(section: StructuredSection) -> str:
    if section.heading and section.content and section.content != section.heading:
        return f"{section.heading}\n{section.content}".strip()
    return (section.content or section.heading or "").strip()


def _guess_media_type(suffix: str) -> str | None:
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".html": "text/html",
        ".htm": "text/html",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
    }.get(suffix)
