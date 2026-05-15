from __future__ import annotations

from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document as DocxDocument

from app.core.exceptions import ValidationError
from app.db.enums import Criticality, DocumentType, SourceStatus, SourceSyncMode, SourceType
from app.domain.services.knowledge_core import KnowledgeSourceService, KnowledgeUpdateService
from app.integrations.knowledge.content_loader import normalize_document_payload
from app.integrations.knowledge.knowledge_extraction import (
    DocumentMemoryLlmConfig,
    extract_document_memory,
)
from app.integrations.knowledge.source_security import validate_source_base_uri


class _ProcessingResultsRepo:
    def __init__(
        self,
        latest_success_by_source: dict[str, object] | None = None,
        latest_by_source: dict[str, object] | None = None,
    ) -> None:
        self.latest_success_by_source = latest_success_by_source or {}
        self.latest_by_source = latest_by_source or {}

    def get_latest_for_source(self, source_id: str):
        return self.latest_by_source.get(source_id)

    def get_latest_success_for_source(self, source_id: str):
        return self.latest_success_by_source.get(source_id)


class _DocumentsRepo:
    def list_for_source(self, _source_id: str, include_archived: bool = True):
        return []


def test_validate_url_source_rejects_local_path_and_file_uri(tmp_path: Path) -> None:
    with pytest.raises(ValidationError):
        validate_source_base_uri(source_type=SourceType.URL, base_uri=str(tmp_path))
    with pytest.raises(ValidationError):
        validate_source_base_uri(source_type=SourceType.URL, base_uri=tmp_path.as_uri())


def test_serialize_source_uses_source_level_last_sync_and_public_type() -> None:
    now = datetime.now(UTC)
    service = KnowledgeSourceService.__new__(KnowledgeSourceService)
    service.processing_results = _ProcessingResultsRepo(
        latest_success_by_source={"src-1": SimpleNamespace(processed_at=now - timedelta(days=10))},
        latest_by_source={
            "src-1": SimpleNamespace(processed_at=now - timedelta(days=2), error_code=None)
        },
    )
    service.documents = _DocumentsRepo()
    service.settings = SimpleNamespace(knowledge_auto_sync_interval_days=30)

    source = SimpleNamespace(
        source_id="src-1",
        knowledge_base_id="kb-1",
        source_type=SourceType.URL_LIST,
        name="Portal",
        base_uri="https://example.com/index.html",
        criticality=Criticality.OPTIONAL,
        status=SourceStatus.ACTIVE,
        refresh_policy="monthly",
        sync_mode=SourceSyncMode.FULL_SCAN,
        source_metadata={},
        created_at=now - timedelta(days=20),
        last_discovered_at=now - timedelta(days=12),
    )

    payload = KnowledgeSourceService._serialize_source(service, source)

    assert payload["source_type"] == SourceType.URL
    assert payload["last_sync_time"] == now - timedelta(days=10)
    assert payload["next_sync_time"] == now + timedelta(days=20)


def test_run_due_scheduled_syncs_targets_only_due_auto_sources() -> None:
    service = KnowledgeUpdateService.__new__(KnowledgeUpdateService)
    now = datetime.now(UTC)
    base = SimpleNamespace(knowledge_base_id="kb-1", status=SourceStatus.ACTIVE)
    due_source = SimpleNamespace(
        source_id="src-due",
        refresh_policy="monthly",
        last_discovered_at=now - timedelta(days=45),
        created_at=now - timedelta(days=60),
    )
    fresh_source = SimpleNamespace(
        source_id="src-fresh",
        refresh_policy="monthly",
        last_discovered_at=now - timedelta(days=2),
        created_at=now - timedelta(days=60),
    )
    manual_source = SimpleNamespace(
        source_id="src-manual",
        refresh_policy="manual",
        last_discovered_at=now - timedelta(days=90),
        created_at=now - timedelta(days=90),
    )
    service.settings = SimpleNamespace(knowledge_auto_sync_interval_days=30)
    service.sources = SimpleNamespace(
        list_active=lambda knowledge_base_id=None: [due_source, fresh_source, manual_source]
    )
    service.processing_results = _ProcessingResultsRepo(
        latest_success_by_source={
            "src-fresh": SimpleNamespace(processed_at=now - timedelta(days=1))
        }
    )
    created_payloads: list[object] = []

    def _create_run(**kwargs):
        created_payloads.append(kwargs["payload"])
        return SimpleNamespace(update_run_id="run-1", knowledge_base_id="kb-1")

    service._create_run = _create_run
    service.session = None

    import app.domain.services.knowledge_core as knowledge_core_module

    original_service = knowledge_core_module.KnowledgeBaseService
    knowledge_core_module.KnowledgeBaseService = lambda session: SimpleNamespace(
        bases=SimpleNamespace(list_visible=lambda: [base])
    )
    try:
        payload = KnowledgeUpdateService.run_due_scheduled_syncs(
            service, now=now, execute_inline=False
        )
    finally:
        knowledge_core_module.KnowledgeBaseService = original_service

    assert payload["started_knowledge_base_ids"] == ["kb-1"]
    assert created_payloads[0].source_scope.value == "selected"
    assert created_payloads[0].selected_source_ids == ["src-due"]


def test_markdown_and_docx_source_locations_are_more_precise() -> None:
    markdown = b"# Intro\n\nOverview line\n\n## Target\n\nDetailed constraint"
    md_payload = normalize_document_payload("guide.md", markdown)
    assert md_payload.sections[0].source_location.startswith("lines:")
    assert "heading:intro" in md_payload.sections[0].source_location

    document = DocxDocument()
    document.add_heading("Business Architecture", level=1)
    document.add_paragraph("Application component shall expose API.")
    buffer = BytesIO()
    document.save(buffer)
    docx_payload = normalize_document_payload("solution.docx", buffer.getvalue())
    assert docx_payload.sections[0].source_location.startswith("docx:block:")
    assert "heading:business-architecture" in docx_payload.sections[0].source_location


def test_unsupported_normalization_format_uses_text_fallback() -> None:
    payload = normalize_document_payload("notes.csv", b"hello")

    assert payload.parser_name == "text-fallback"
    assert payload.metadata["fallback_used"] is True


def test_document_memory_reports_llm_fallback_diagnostics() -> None:
    memory = extract_document_memory(
        document_title="API Guide",
        document_type=DocumentType.API,
        normalized_text="GET /status shall be available. Risk of timeout is noted.",
        chunks=[
            {
                "document_chunk_id": "chunk-1",
                "title": "API",
                "content": "GET /status shall be available. Risk of timeout is noted.",
                "source_location": "lines:1-1",
            }
        ],
        llm_config=DocumentMemoryLlmConfig(
            provider="openai",
            base_url="http://127.0.0.1:1/v1",
            model_id="gpt-test",
            timeout_sec=0.01,
        ),
    )

    assert memory.llm_attempted is True
    assert memory.fallback_applied is True
    assert memory.extraction_method == "heuristic"
    assert memory.fallback_reason
    assert any(
        item.structured_payload.get("extraction_method") == "heuristic" for item in memory.items
    )
