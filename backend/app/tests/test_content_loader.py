from __future__ import annotations

import base64
import io

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.integrations.knowledge import content_loader
from app.integrations.knowledge.content_loader import normalize_document_payload

PDF_SAMPLE_BASE64 = "JVBERi0xLjMKJZOMi54gUmVwb3J0TGFiIEdlbmVyYXRlZCBQREYgZG9jdW1lbnQgKG9wZW5zb3VyY2UpCjEgMCBvYmoKPDwKL0YxIDIgMCBSCj4+CmVuZG9iagoyIDAgb2JqCjw8Ci9CYXNlRm9udCAvSGVsdmV0aWNhIC9FbmNvZGluZyAvV2luQW5zaUVuY29kaW5nIC9OYW1lIC9GMSAvU3VidHlwZSAvVHlwZTEgL1R5cGUgL0ZvbnQKPj4KZW5kb2JqCjMgMCBvYmoKPDwKL0NvbnRlbnRzIDcgMCBSIC9NZWRpYUJveCBbIDAgMCA1OTUuMjc1NiA4NDEuODg5OCBdIC9QYXJlbnQgNiAwIFIgL1Jlc291cmNlcyA8PAovRm9udCAxIDAgUiAvUHJvY1NldCBbIC9QREYgL1RleHQgL0ltYWdlQiAvSW1hZ2VDIC9JbWFnZUkgXQo+PiAvUm90YXRlIDAgL1RyYW5zIDw8Cgo+PiAKICAvVHlwZSAvUGFnZQo+PgplbmRvYmoKNCAwIG9iago8PAovUGFnZU1vZGUgL1VzZU5vbmUgL1BhZ2VzIDYgMCBSIC9UeXBlIC9DYXRhbG9nCj4+CmVuZG9iago1IDAgb2JqCjw8Ci9BdXRob3IgKGFub255bW91cykgL0NyZWF0aW9uRGF0ZSAoRDoyMDI2MDQwMzA1Mzk1NiswMCcwMCcpIC9DcmVhdG9yIChhbm9ueW1vdXMpIC9LZXl3b3JkcyAoKSAvTW9kRGF0ZSAoRDoyMDI2MDQwMzA1Mzk1NiswMCcwMCcpIC9Qcm9kdWNlciAoUmVwb3J0TGFiIFBERiBMaWJyYXJ5IC0gXChvcGVuc291cmNlXCkpIAogIC9TdWJqZWN0ICh1bnNwZWNpZmllZCkgL1RpdGxlICh1bnRpdGxlZCkgL1RyYXBwZWQgL0ZhbHNlCj4+CmVuZG9iago2IDAgb2JqCjw8Ci9Db3VudCAxIC9LaWRzIFsgMyAwIFIgXSAvVHlwZSAvUGFnZXMKPj4KZW5kb2JqCjcgMCBvYmoKPDwKL0ZpbHRlciBbIC9BU0NJSTg1RGVjb2RlIC9GbGF0ZURlY29kZSBdIC9MZW5ndGggMTU3Cj4+CnN0cmVhbQpHYXJXb1ltUz81Jjs5Iis6W3NjZCZUOU85KlJWMipbK0gwY2ledEMyT2g5aUspcTtfayMsVzciJ0FPRk1HTE9IIVI8cTQ4a1FORishXjNkOkN1Ii0hKCRBXSYqODMxdEEoLFVSMjheWFdGZDRNIjJyaDo8O3RQc21ybUJoX2w+XjBXO0huaT9EOzptNE1iY0Z0NUV0TDxkKj0jZ34+ZW5kc3RyZWFtCmVuZG9iagp4cmVmCjAgOAowMDAwMDAwMDAwIDY1NTM1IGYgCjAwMDAwMDAwNjEgMDAwMDAgbiAKMDAwMDAwMDA5MiAwMDAwMCBuIAowMDAwMDAwMTk5IDAwMDAwIG4gCjAwMDAwMDA0MDIgMDAwMDAgbiAKMDAwMDAwMDQ3MCAwMDAwMCBuIAowMDAwMDAwNzMxIDAwMDAwIG4gCjAwMDAwMDA3OTAgMDAwMDAgbiAKdHJhaWxlcgo8PAovSUQgCls8YTE3ZTM3N2ZiMWJjNjk2YzljZTMxMjVjNjc2NzBlNzM+PGExN2UzNzdmYjFiYzY5NmM5Y2UzMTI1YzY3NjcwZTczPl0KJSBSZXBvcnRMYWIgZ2VuZXJhdGVkIFBERiBkb2N1bWVudCAtLSBkaWdlc3QgKG9wZW5zb3VyY2UpCgovSW5mbyA1IDAgUgovUm9vdCA0IDAgUgovU2l6ZSA4Cj4+CnN0YXJ0eHJlZgoxMDM3CiUlRU9GCg=="


def test_normalize_pdf_payload_extracts_text_and_sections() -> None:
    payload = normalize_document_payload("sample.pdf", base64.b64decode(PDF_SAMPLE_BASE64))
    assert payload.content_format == "pdf"
    assert payload.parser_name == "pypdf"
    assert "Architecture baseline PDF" in payload.text
    assert payload.metadata["page_count"] == 1
    assert payload.sections[0].source_location == "page:1"


def test_normalize_pdf_payload_skips_duplicate_and_legal_pages(monkeypatch) -> None:
    class _Page:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class _Reader:
        def __init__(self, _handle) -> None:
            self.pages = [
                _Page("Legal disclaimer. Confidential. Copyright. All rights reserved."),
                _Page("Architecture component exposes API for order processing."),
                _Page("Architecture component exposes API for order processing."),
            ]

    monkeypatch.setattr(content_loader, "PdfReader", _Reader)

    payload = normalize_document_payload("guide.pdf", b"%PDF")

    assert payload.text == "Page 2\nArchitecture component exposes API for order processing."
    assert payload.metadata["skipped_page_count"] == 2
    assert payload.metadata["skipped_pages"] == [
        {"page_number": 1, "reason": "legal_disclaimer"},
        {"page_number": 3, "reason": "duplicate_page"},
    ]


def test_normalize_pdf_payload_skips_toc_and_keeps_content_with_legal_footer(monkeypatch) -> None:
    class _Page:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    content_with_footer = (
        "Architecture governance defines the architecture board, compliance review, "
        "stakeholder concerns, business capability map, target architecture, baseline "
        "architecture, transition architecture, and implementation governance. "
        * 4
    ) + "© The Open Group, All Rights Reserved. Evaluation Copy. Not for redistribution"

    class _Reader:
        def __init__(self, _handle) -> None:
            self.pages = [
                _Page("TOGAF Standard Evaluation Copy. © The Open Group, All Rights Reserved. Not for redistribution"),
                _Page(
                    "Table of Contents\n"
                    "1. Architecture Vision . . . . . . . . . . . . . . 1\n"
                    "2. Business Architecture . . . . . . . . . . . . 5\n"
                    "3. Data Architecture . . . . . . . . . . . . . . 9\n"
                    "4. Application Architecture . . . . . . . . . 12\n"
                    "5. Technology Architecture . . . . . . . . . 18\n"
                ),
                _Page(
                    "6. Opportunities and Solutions . . . . . . . 24\n"
                    "7. Migration Planning . . . . . . . . . . . . 31\n"
                    "8. Implementation Governance . . . . . . . 40\n"
                    "9. Architecture Change Management . . . . 49\n"
                    "10. Requirements Management . . . . . . . 55\n"
                    "11. Architecture Repository . . . . . . . . . 63\n"
                    "12. Architecture Board . . . . . . . . . . . 70\n"
                    "13. Architecture Compliance . . . . . . . . 78\n"
                ),
                _Page(content_with_footer),
            ]

    monkeypatch.setattr(content_loader, "PdfReader", _Reader)

    payload = normalize_document_payload("togaf.pdf", b"%PDF")

    assert "Architecture governance defines the architecture board" in payload.text
    assert "Table of Contents" not in payload.text
    assert payload.metadata["skipped_pages"] == [
        {"page_number": 1, "reason": "legal_disclaimer"},
        {"page_number": 2, "reason": "table_of_contents"},
        {"page_number": 3, "reason": "table_of_contents"},
    ]


def test_normalize_pdf_payload_falls_back_when_parser_requires_optional_crypto(
    monkeypatch,
) -> None:
    def _raise_optional_crypto_error(_handle):
        raise RuntimeError("cryptography>=3.1 is required for AES algorithm")

    monkeypatch.setattr(content_loader, "PdfReader", _raise_optional_crypto_error)

    payload = normalize_document_payload("catalog.pdf", b"\x80\x81\x00\x00")

    assert payload.parser_name == "binary-fallback"
    assert payload.metadata["fallback_used"] is True
    assert payload.metadata["pdf_fallback_reason"] == "parse_failed"
    assert "cryptography>=3.1 is required for AES algorithm" in payload.metadata["pdf_parse_error"]


def test_normalize_docx_payload_preserves_headings() -> None:
    document = DocxDocument()
    document.add_heading("Business Architecture", level=1)
    document.add_paragraph("Application component shall expose API.")
    buffer = io.BytesIO()
    document.save(buffer)
    payload = normalize_document_payload("solution.docx", buffer.getvalue())
    assert payload.content_format == "docx"
    assert payload.sections[0].heading == "Business Architecture"
    assert "Application component shall expose API." in payload.text


def test_normalize_xlsx_payload_extracts_sheets_and_rows() -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "SID Entities"
    worksheet.append(["Entity", "Description"])
    worksheet.append(["Customer", "Represents a party using services"])
    worksheet.append([None, None])
    worksheet.append(["ProductOffering", "Catalog item offered to customers"])
    second_sheet = workbook.create_sheet("Empty Sheet")

    buffer = io.BytesIO()
    workbook.save(buffer)
    workbook.close()

    payload = normalize_document_payload("sid.xlsx", buffer.getvalue())

    assert payload.content_format == "xlsx"
    assert payload.parser_name == "openpyxl"
    assert payload.metadata["sheet_count"] == 2
    assert payload.metadata["section_count"] == 1
    assert payload.metadata["non_empty_row_count"] == 3
    assert payload.sections[0].heading == "Sheet: SID Entities"
    assert payload.sections[0].source_location.startswith("xlsx:sheet:1;rows:1-4")
    assert "Customer | Represents a party using services" in payload.text
    assert second_sheet.title == "Empty Sheet"


def test_normalize_xlsx_payload_skips_duplicate_summary_sheets() -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "All Domains"
    worksheet.append(["Duplicated aggregate row"])
    for title in ("Product", "Service", "Resource"):
        domain_sheet = workbook.create_sheet(title)
        domain_sheet.append(["Entity", "Description"])
        domain_sheet.append([title, f"{title} domain details"])

    buffer = io.BytesIO()
    workbook.save(buffer)
    workbook.close()

    payload = normalize_document_payload("sid.xlsx", buffer.getvalue())

    assert payload.metadata["sheet_count"] == 4
    assert payload.metadata["section_count"] == 3
    assert payload.metadata["skipped_sheets"][0]["sheet_name"] == "All Domains"
    assert "Duplicated aggregate row" not in payload.text
    assert "Product domain details" in payload.text


def test_normalize_html_payload_extracts_headings() -> None:
    html = b"<html><body><h1>Technology Architecture</h1><p>Redis is required.</p></body></html>"
    payload = normalize_document_payload("guide.html", html)
    assert payload.content_format == "html"
    assert payload.sections[0].heading == "Technology Architecture"
    assert "Redis is required." in payload.text


def test_normalize_markdown_payload_splits_sections() -> None:
    markdown = b"# General Information\n\nSystem overview.\n\n## Business Architecture\n\nBusiness actor and process."
    payload = normalize_document_payload("guide.md", markdown)
    assert payload.content_format == "markdown"
    assert len(payload.sections) == 2
    assert payload.sections[0].heading == "General Information"
    assert payload.sections[1].heading == "Business Architecture"


def test_normalize_html_payload_uses_precise_source_location() -> None:
    html = b"<html><body><h2>Apps</h2><p>Main flow.</p></body></html>"
    payload = normalize_document_payload("apps.html", html)
    assert payload.sections[0].source_location.startswith("html:block:")
    assert "heading:apps" in payload.sections[0].source_location


def test_normalize_unknown_extension_uses_text_fallback() -> None:
    payload = normalize_document_payload("catalog.csv", b"sku,name\n1,Router")
    assert payload.parser_name == "text-fallback"
    assert payload.metadata["fallback_used"] is True
    assert "sku,name" in payload.text


def test_normalize_binary_unknown_extension_uses_binary_fallback() -> None:
    payload = normalize_document_payload("archive.bin", b"\x50\x4b\x03\x04\x00\x00\x01\x02")
    assert payload.parser_name == "binary-fallback"
    assert payload.metadata["fallback_used"] is True
    assert "Binary document fallback" in payload.text
